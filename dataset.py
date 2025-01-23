import docx
from docx import Document
import re
import json

# Đọc nội dung từ tệp Word
file_path = 'dataset/36_2024_QH15_444251.docx'
document = docx.Document()
document = Document(file_path)

content = ""
for paragraph in document.paragraphs:
    content += paragraph.text + "\n"

# Tách các chương
chuong_pattern = re.compile(r'Chương\\s+(\\w+)\\s+(.*?)\\n')
dieu_pattern = re.compile(r'Điều\\s+(\\d+)\\.\\s+(.*?)\\n(.*?)(?=(\\nĐiều\\s+\\d+|\\nChương\\s+|$))', re.S)

result = []

for chuong_match in chuong_pattern.finditer(content):
    chuong_number, chuong_title = chuong_match.groups()
    chuong_start = chuong_match.end()
    
    next_chuong = chuong_pattern.search(content, chuong_start)
    chuong_end = next_chuong.start() if next_chuong else len(content)

    chuong_content = content[chuong_start:chuong_end]

    # Tách các điều trong chương
    dieu_matches = dieu_pattern.finditer(chuong_content)
    dieu_list = []
    for dieu_match in dieu_matches:
        dieu_number, dieu_title, dieu_content, _ = dieu_match.groups()
        dieu_list.append({
            "dieu": int(dieu_number),
            "title": dieu_title.strip(),
            "content": dieu_content.strip()
        })

    result.append({
        "chuong": {
            "number": chuong_number,
            "title": chuong_title.strip(),
            "dieu": dieu_list
        }
    })

# Xuất dữ liệu ra file JSON
output_path = '/mnt/data/luat_hon_nhan_va_gia_dinh.json'
with open(output_path, 'w', encoding='utf-8') as json_file:
    json.dump(result, json_file, ensure_ascii=False, indent=4)

print(f"Dữ liệu đã được xử lý và lưu tại: {output_path}")
